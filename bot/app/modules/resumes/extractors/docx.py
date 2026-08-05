from io import BytesIO
from zipfile import BadZipFile, ZipFile, is_zipfile

from docx import Document

from app.core.exceptions import BotError
from app.modules.resumes.extractors.base import ExtractedDocument


class DocxResumeExtractor:
    def __init__(self, max_uncompressed_bytes: int) -> None:
        self.max_uncompressed_bytes = max_uncompressed_bytes

    def extract(self, content: bytes) -> ExtractedDocument:
        self._validate_archive(content)

        try:
            document = Document(BytesIO(content))
        except Exception as exception:
            raise BotError("INVALID_DOCX", "The DOCX file is invalid.", 422) from exception

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        table_rows = [
            " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for table in document.tables
            for row in table.rows
        ]
        full_text = "\n".join(value for value in [*paragraphs, *table_rows] if value).strip()

        if not full_text:
            raise BotError("EMPTY_DOCUMENT", "No readable text was found in the resume.", 422)

        return ExtractedDocument(
            full_text=full_text,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={
                "paragraph_count": len([value for value in paragraphs if value]),
                "table_count": len(document.tables),
            },
        )

    def _validate_archive(self, content: bytes) -> None:
        if not is_zipfile(BytesIO(content)):
            raise BotError("INVALID_DOCX", "The DOCX file is invalid.", 422)

        try:
            with ZipFile(BytesIO(content)) as archive:
                names = set(archive.namelist())
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise BotError("INVALID_DOCX", "The DOCX file is invalid.", 422)

                uncompressed_size = sum(entry.file_size for entry in archive.infolist())
                if uncompressed_size > self.max_uncompressed_bytes:
                    raise BotError("DOCX_SIZE_LIMIT", "The DOCX content is too large.", 422)
        except BadZipFile as exception:
            raise BotError("INVALID_DOCX", "The DOCX file is invalid.", 422) from exception
