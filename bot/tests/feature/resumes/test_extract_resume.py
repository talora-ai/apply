import time
from collections.abc import Callable
from io import BytesIO
from uuid import uuid4

import pymupdf
from docx import Document
from fastapi.testclient import TestClient


def create_pdf(text: str = "Gustavo Martim\nHabilidades\nPHP\nLaravel") -> bytes:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    content = document.tobytes()
    document.close()
    return content


def create_docx() -> bytes:
    output = BytesIO()
    document = Document()
    document.add_paragraph("Gustavo Martim")
    document.add_paragraph("Experiência")
    document.add_paragraph("Backend Developer — Talora | 2024\u2013Atual")
    document.add_paragraph("Habilidades")
    document.add_paragraph("PHP, Laravel, PostgreSQL")
    document.add_paragraph("Projetos selecionados")
    document.add_paragraph("Talora Apply")
    document.add_paragraph("• Plataforma de apoio a candidaturas.")
    document.add_paragraph("• Backend Laravel e frontend Next.js.")
    document.add_paragraph("Formação e Idiomas")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Engenharia de Software"
    table.cell(0, 1).text = "Português: Nativo\nInglês: Leitura técnica"
    document.save(output)
    return output.getvalue()


def test_requires_service_authentication(client: TestClient) -> None:
    response = client.post(
        "/api/v1/resumes/extract",
        files={"file": ("resume.pdf", create_pdf(), "application/pdf")},
    )

    assert response.status_code == 401
    assert response.json()["error"]["code"] == "UNAUTHORIZED"


def test_rejects_an_invalid_service_token(client: TestClient) -> None:
    response = client.post(
        "/api/v1/resumes/extract",
        headers={"Authorization": "Bearer invalid-service-token"},
        files={"file": ("resume.pdf", create_pdf(), "application/pdf")},
    )

    assert response.status_code == 401
    assert response.json()["error"]["code"] == "UNAUTHORIZED"


def test_extracts_a_pdf_resume(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    content = create_pdf()
    headers = auth_headers(content)
    response = client.post(
        "/api/v1/resumes/extract",
        headers=headers,
        files={"file": ("resume.pdf", content, "application/pdf")},
    )

    assert response.status_code == 200
    assert response.headers["cache-control"] == "no-store"
    payload = response.json()
    assert payload["schema_version"] == "1.4"
    assert payload["processing_id"] == headers["X-Talora-Processing-Id"]
    assert payload["document"]["sha256"] == headers["X-Talora-Content-SHA256"]
    assert payload["document"]["mime_type"] == "application/pdf"
    assert payload["document"]["page_count"] == 1
    assert payload["document"]["ats"]["ats_friendly"] is True
    assert payload["document"]["ats"]["layout_type"] == "single_column"
    assert "Gustavo Martim" in payload["content"]["full_text"]
    assert payload["content"]["sections"]["skills"] == ["PHP", "Laravel"]


def test_extracts_a_docx_resume(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    content = create_docx()
    response = client.post(
        "/api/v1/resumes/extract",
        headers=auth_headers(content),
        files={
            "file": (
                "resume.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["schema_version"] == "1.4"
    assert "Backend Developer — Talora" in payload["content"]["full_text"]
    sections = payload["content"]["sections"]
    assert payload["document"]["ats"]["ats_friendly"] is True
    assert payload["document"]["ats"]["layout_type"] == "document_flow"
    assert sections["experiences"] == [
        {
            "position": "Backend Developer",
            "company": "Talora",
            "period": "2024\u2013Atual",
            "start_date": "2024",
            "end_date": None,
            "is_current": True,
            "description": [],
        }
    ]
    assert sections["skills"] == ["PHP", "Laravel", "PostgreSQL", "Next.js"]
    assert sections["projects"] == [
        {
            "name": "Talora Apply",
            "description": [
                "Plataforma de apoio a candidaturas.",
                "Backend Laravel e frontend Next.js.",
            ],
        }
    ]
    assert sections["education"] == ["Engenharia de Software"]
    assert sections["languages"] == [
        {"name": "Português", "proficiency": "Nativo"},
        {"name": "Inglês", "proficiency": "Leitura técnica"},
    ]


def test_rejects_an_empty_file(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    response = client.post(
        "/api/v1/resumes/extract",
        headers=auth_headers(b""),
        files={"file": ("resume.pdf", b"", "application/pdf")},
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "EMPTY_FILE"


def test_rejects_an_unsupported_file(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    content = b"plain text"
    response = client.post(
        "/api/v1/resumes/extract",
        headers=auth_headers(content),
        files={"file": ("resume.txt", content, "text/plain")},
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "UNSUPPORTED_FILE_TYPE"


def test_rejects_a_file_over_10_mb(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    content = b"%PDF-" + b"0" * (10 * 1024 * 1024)
    response = client.post(
        "/api/v1/resumes/extract",
        headers=auth_headers(content),
        files={"file": ("resume.pdf", content, "application/pdf")},
    )

    assert response.status_code == 413
    assert response.json()["error"]["code"] == "FILE_SIZE_LIMIT"


def test_rejects_a_mismatched_extension(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    content = create_pdf()
    response = client.post(
        "/api/v1/resumes/extract",
        headers=auth_headers(content),
        files={"file": ("resume.docx", content, "application/pdf")},
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "UNSUPPORTED_FILE_TYPE"


def test_rejects_a_tampered_document(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    original = create_pdf("original")
    tampered = create_pdf("tampered")

    response = client.post(
        "/api/v1/resumes/extract",
        headers=auth_headers(original),
        files={"file": ("resume.pdf", tampered, "application/pdf")},
    )

    assert response.status_code == 401
    assert response.json()["error"]["code"] == "CONTENT_INTEGRITY_FAILED"


def test_rejects_an_expired_signature(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    content = create_pdf()

    response = client.post(
        "/api/v1/resumes/extract",
        headers=auth_headers(content, timestamp=int(time.time()) - 120),
        files={"file": ("resume.pdf", content, "application/pdf")},
    )

    assert response.status_code == 401
    assert response.json()["error"]["code"] == "EXPIRED_SIGNATURE"


def test_rejects_a_replayed_request(
    client: TestClient,
    auth_headers: Callable[..., dict[str, str]],
) -> None:
    content = create_pdf()
    headers = auth_headers(content, processing_id=str(uuid4()), nonce=str(uuid4()))

    first = client.post(
        "/api/v1/resumes/extract",
        headers=headers,
        files={"file": ("resume.pdf", content, "application/pdf")},
    )
    replay = client.post(
        "/api/v1/resumes/extract",
        headers=headers,
        files={"file": ("resume.pdf", content, "application/pdf")},
    )

    assert first.status_code == 200
    assert replay.status_code == 409
    assert replay.json()["error"]["code"] == "REPLAY_DETECTED"
